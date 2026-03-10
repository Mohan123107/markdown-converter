#!/usr/bin/env python3
"""Convert Beacon Nudge files (.docx, .xlsx, .md) to Markdown."""

import os
import shutil
from pathlib import Path

from docx import Document
from openpyxl import load_workbook

SOURCE_DIR = Path(__file__).parent / "Beacon Nudge"
OUTPUT_DIR = Path(__file__).parent / "converted_md"

SKIP_EXTENSIONS = {".png", ".zip"}


def escape_md_table_cell(text: str) -> str:
    """Escape pipe characters in markdown table cells."""
    return text.replace("|", "\\|").strip()


def convert_docx_to_md(docx_path: Path) -> str:
    """Convert a .docx file to markdown text."""
    doc = Document(str(docx_path))
    lines = []

    for element in doc.element.body:
        tag = element.tag.split("}")[-1]

        if tag == "p":
            # It's a paragraph
            para = None
            for p in doc.paragraphs:
                if p._element is element:
                    para = p
                    break
            if para is None:
                continue

            style_name = para.style.name if para.style else ""
            text = _get_paragraph_md(para)

            if not text.strip():
                lines.append("")
                continue

            if style_name.startswith("Heading"):
                try:
                    level = int(style_name.split()[-1])
                except (ValueError, IndexError):
                    level = 1
                level = min(level, 6)
                lines.append(f"{'#' * level} {text}")
            elif style_name in ("List Bullet", "List Bullet 2", "List Bullet 3"):
                indent = 0
                if "2" in style_name:
                    indent = 1
                elif "3" in style_name:
                    indent = 2
                lines.append(f"{'  ' * indent}- {text}")
            elif style_name in ("List Number", "List Number 2", "List Number 3"):
                indent = 0
                if "2" in style_name:
                    indent = 1
                elif "3" in style_name:
                    indent = 2
                lines.append(f"{'  ' * indent}1. {text}")
            else:
                lines.append(text)

        elif tag == "tbl":
            # It's a table
            tbl = None
            for t in doc.tables:
                if t._element is element:
                    tbl = t
                    break
            if tbl is None:
                continue

            lines.append("")
            for i, row in enumerate(tbl.rows):
                cells = [escape_md_table_cell(cell.text) for cell in row.cells]
                lines.append("| " + " | ".join(cells) + " |")
                if i == 0:
                    lines.append("| " + " | ".join("---" for _ in cells) + " |")
            lines.append("")

    return "\n".join(lines) + "\n"


def _get_paragraph_md(para) -> str:
    """Extract paragraph text with bold/italic formatting."""
    parts = []
    for run in para.runs:
        text = run.text
        if not text:
            continue
        if run.bold and run.italic:
            text = f"***{text}***"
        elif run.bold:
            text = f"**{text}**"
        elif run.italic:
            text = f"*{text}*"
        parts.append(text)
    return "".join(parts)


def convert_xlsx_to_md(xlsx_path: Path) -> str:
    """Convert an .xlsx file to markdown (each sheet as a table)."""
    wb = load_workbook(str(xlsx_path), read_only=True, data_only=True)
    sections = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = []
        for row in ws.iter_rows(values_only=True):
            # Convert each cell to string
            cells = [escape_md_table_cell(str(c)) if c is not None else "" for c in row]
            rows.append(cells)

        if not rows:
            continue

        lines = [f"## Sheet: {sheet_name}", ""]

        # Normalize column count
        max_cols = max(len(r) for r in rows)
        for r in rows:
            while len(r) < max_cols:
                r.append("")

        # First row as header
        lines.append("| " + " | ".join(rows[0]) + " |")
        lines.append("| " + " | ".join("---" for _ in rows[0]) + " |")

        for row in rows[1:]:
            lines.append("| " + " | ".join(row) + " |")

        sections.append("\n".join(lines))

    wb.close()
    return "\n\n".join(sections) + "\n"


def main():
    stats = {"docx": 0, "xlsx": 0, "md_copied": 0, "skipped": 0}

    if OUTPUT_DIR.exists():
        shutil.rmtree(OUTPUT_DIR)
    OUTPUT_DIR.mkdir(parents=True)

    for root, _dirs, files in os.walk(SOURCE_DIR):
        root_path = Path(root)
        rel_dir = root_path.relative_to(SOURCE_DIR)
        out_dir = OUTPUT_DIR / rel_dir
        out_dir.mkdir(parents=True, exist_ok=True)

        for filename in files:
            src_file = root_path / filename
            ext = src_file.suffix.lower()

            if ext in SKIP_EXTENSIONS:
                print(f"  SKIP  {src_file.relative_to(SOURCE_DIR)}")
                stats["skipped"] += 1
                continue

            if ext == ".md":
                dest = out_dir / filename
                shutil.copy2(src_file, dest)
                print(f"  COPY  {src_file.relative_to(SOURCE_DIR)} -> {dest.relative_to(OUTPUT_DIR)}")
                stats["md_copied"] += 1

            elif ext == ".docx":
                md_name = src_file.stem + ".md"
                dest = out_dir / md_name
                md_text = convert_docx_to_md(src_file)
                dest.write_text(md_text, encoding="utf-8")
                print(f"  DOCX  {src_file.relative_to(SOURCE_DIR)} -> {dest.relative_to(OUTPUT_DIR)}")
                stats["docx"] += 1

            elif ext == ".xlsx":
                md_name = src_file.stem + ".md"
                dest = out_dir / md_name
                md_text = convert_xlsx_to_md(src_file)
                dest.write_text(md_text, encoding="utf-8")
                print(f"  XLSX  {src_file.relative_to(SOURCE_DIR)} -> {dest.relative_to(OUTPUT_DIR)}")
                stats["xlsx"] += 1

            else:
                print(f"  SKIP  {src_file.relative_to(SOURCE_DIR)} (unknown type)")
                stats["skipped"] += 1

    print("\n--- Summary ---")
    print(f"  DOCX converted: {stats['docx']}")
    print(f"  XLSX converted: {stats['xlsx']}")
    print(f"  MD copied:      {stats['md_copied']}")
    print(f"  Skipped:        {stats['skipped']}")
    print(f"  Total output:   {stats['docx'] + stats['xlsx'] + stats['md_copied']} .md files")


if __name__ == "__main__":
    main()
